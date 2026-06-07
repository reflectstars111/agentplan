"""WordParser — .docx file parsing for Agent-OS.

Uses python-docx to extract structured text from Word documents.
Maps to agent_os_initial_plan.md §10.1 (Word input).
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from src.models.chunk import DocumentChunk, ChunkType, ChunkLocation, TrustLevel


@dataclass
class WordParseResult:
    """Result of Word document parsing."""
    chunks: list[DocumentChunk] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    parser_used: str = "none"


class WordParser:
    """Parse .docx files into DocumentChunks using python-docx."""

    def parse(self, file_path: Path, source_id: str) -> WordParseResult:
        """Parse a .docx file into structured chunks.

        Extracts paragraphs, headings (by style), and tables.
        Falls back gracefully if python-docx is not installed.
        """
        try:
            from docx import Document
        except ImportError:
            return WordParseResult(
                chunks=[],
                metadata={"error": "python-docx not installed. pip install python-docx",
                          "parser": "none"},
                parser_used="none",
            )

        try:
            doc = Document(str(file_path))
        except Exception as e:
            return WordParseResult(
                chunks=[],
                metadata={"error": str(e), "parser": "none"},
                parser_used="none",
            )

        chunks = []
        idx = 0
        pages_estimate = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect heading by style name
            is_heading = para.style.name.startswith("Heading") if para.style else False
            chunk_type = ChunkType.HEADING if is_heading else ChunkType.PARAGRAPH

            chunks.append(DocumentChunk(
                chunk_id=f"chunk_{source_id}_{idx:04d}",
                source_id=source_id,
                source_type="word",
                text=text,
                chunk_type=chunk_type,
                trust_level=TrustLevel.USER_PROVIDED_DATA,
            ))
            idx += 1

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                chunks.append(DocumentChunk(
                    chunk_id=f"chunk_{source_id}_t{table_idx:04d}",
                    source_id=source_id,
                    source_type="word",
                    text=table_text,
                    chunk_type=ChunkType.TABLE,
                    trust_level=TrustLevel.USER_PROVIDED_DATA,
                ))

        return WordParseResult(
            chunks=chunks,
            metadata={
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "parser": "python-docx",
            },
            parser_used="python-docx",
        )

    @staticmethod
    def is_available() -> bool:
        """Check if python-docx is installed."""
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            return False
